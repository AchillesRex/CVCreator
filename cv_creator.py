import tkinter as tk
from tkinter import messagebox
import json
import re
from tkinter import filedialog
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import ttk
import traceback

photo_path = None  # Global variable to store the photo path

def upload_photo(label):
    global photo_path  # Declare global to modify it
    photo_path = filedialog.askopenfilename(filetypes=[("Image Files", "*.png;*.jpg;*.jpeg")])
    if photo_path:
        img = Image.open(photo_path)
        img.thumbnail((100, 100))  # Resize to fit, e.g., 100x100 pixels
        photo = ImageTk.PhotoImage(img)
        label.config(image=photo)
        label.image = photo  # Keep a reference to avoid garbage collection

def remove_photo(label):
    label.config(image='')
    label.image = None

def is_valid_email(email):
    """Validate the email format."""
    email_pattern = re.compile(r"[^@]+@[^@]+\.[^@]+")
    return email_pattern.match(email)

def is_valid_date(date_text):
    """Validate the date format, expecting 'MM/YYYY'."""
    # Adjust the pattern if you expect a different date format
    date_pattern = re.compile(r"\d{2}/\d{4}")
    return date_pattern.match(date_text)

import re

def validate_personal_info(name, email, cellphone, id_number):
    """Validate personal information fields."""
    errors = []
    if not name.strip():
        errors.append("Name cannot be empty.")
    elif len(name) > 50:
        errors.append("Name cannot be more than 50 characters.")
    if not email.strip():
        errors.append("Email cannot be empty.")
    if not cellphone.strip():
        errors.append("Cellphone cannot be empty.")
    if not is_valid_email(email):
        errors.append("Invalid email format.")
    if not (len(cellphone.strip()) == 10 or (cellphone.strip().startswith('+') and len(cellphone.strip()) == 13)):
        errors.append("Cellphone number must be 10 digits (local) or 13 digits (international including '+').")
    if not re.match(r'^\d{13}$', id_number):
        errors.append("ID number must be a 13-digit number.")
    return errors

def add_input_field(root, label_text, row):
    label = tk.Label(root, text=label_text)
    label.grid(row=row, column=0, padx=10, pady=10)
    entry = tk.Entry(root)
    entry.grid(row=row, column=1, padx=10, pady=10)
    return entry

def _on_mousewheel(event, canvas, main_frame):
    if main_frame.winfo_containing(event.x_root, event.y_root):
        canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
def on_main_frame_change(event, main_canvas):
    '''Update the scrollregion whenever the size of the main_frame changes.'''
    main_canvas.configure(scrollregion=main_canvas.bbox("all"))    

def on_frame_configure(canvas):
    '''Reset the scroll region to encompass the inner frame'''
    canvas.configure(scrollregion=canvas.bbox("all"))

def add_labeled_input(parent, label, row, column=0, columnspan=1, **kwargs):
    ttk.Label(parent, text=label, style='TLabel').grid(row=row, column=column, sticky='w', padx=5, pady=5)
    entry = ttk.Entry(parent, **kwargs)
    entry.grid(row=row, column=column+1, columnspan=columnspan, sticky='ew', padx=5, pady=5)
    return entry

def main():
    root = tk.Tk()
    root.title("CV Creator")
    root.geometry("512x768")

    # Define a style
    style = ttk.Style(root)
    style.configure('TLabel', font=('Helvetica', 12))
    style.configure('TButton', font=('Helvetica', 12), borderwidth=1)
    style.configure('TEntry', font=('Helvetica', 12), borderwidth=1)
    style.configure('TFrame', background='white')

    # Create a main canvas and a vertical scrollbar
    main_canvas = tk.Canvas(root)
    main_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    v_scrollbar = tk.Scrollbar(root, orient="vertical", command=main_canvas.yview)
    v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    main_canvas.configure(yscrollcommand=v_scrollbar.set)

    # Create a frame to contain all the widgets
    main_frame = ttk.Frame(root, padding=(10, 10, 10, 10), style='TFrame')
    main_frame.pack(fill=tk.BOTH, expand=True)
    main_canvas.create_window((0, 0), window=main_frame, anchor="nw")
    main_frame.bind('<Configure>', lambda event, canvas=main_canvas: on_main_frame_change(event, canvas))
    root.bind("<MouseWheel>", lambda event: _on_mousewheel(event, main_canvas, main_frame))

    # Define buttons_frame here, before creating buttons
    button_frame = ttk.Frame(main_frame, style='TFrame')
    button_frame.grid(row=13, column=0, columnspan=2, pady=10)  # Adjust grid parameters as needed

    # Initialize the lists for dynamic UI elements
    experience_list = []
    education_list = []
    skills_list = []
    awards_list = []
    certificates_list = []

    # Personal Info Section
    personal_info_label = ttk.Label(main_frame, text="Personal Information", font=('Helvetica', 14, 'bold'))
    personal_info_label.grid(row=0, column=0, columnspan=2, sticky='w')

    personal_info_frame = ttk.Frame(main_frame, padding=(10, 10, 10, 10), style='TFrame')
    personal_info_frame.grid(row=1, column=0, columnspan=2, sticky='ew')
    personal_info_frame.columnconfigure(1, weight=1)  # Make the input fields expand

    name_entry = add_labeled_input(personal_info_frame, "Full Names & Surname:", row=0)
    email_entry = add_labeled_input(personal_info_frame, "Email:", row=1)
    cellphone_entry = add_labeled_input(personal_info_frame, "Cellphone Number:", row=2)
    address_entry = add_labeled_input(personal_info_frame, "Address:", row=3, columnspan=3)
    linkedin_entry = add_labeled_input(personal_info_frame, "LinkedIn Profile:", row=4)
    id_number_entry = add_labeled_input(personal_info_frame, "ID Number:", row=5)
    drivers_license_entry = add_labeled_input(personal_info_frame, "Drivers License Code:", row=6)
    citizenship_entry = add_labeled_input(personal_info_frame, "Citizenship:", row=7)
    passport_number_entry = add_labeled_input(personal_info_frame, "Passport Number:", row=8)

    # Experience Section
    experience_label = ttk.Label(main_frame, text="Experience", font=('Helvetica', 14, 'bold'))
    experience_label.grid(row=4, column=0, columnspan=2, sticky='w', pady=(10, 0))
    experience_frame = ttk.Frame(main_frame, padding=(10, 10, 10, 10), style='TFrame')
    experience_frame.grid(row=5, column=0, columnspan=2, sticky='ew')
    add_experience_button = ttk.Button(experience_frame, text="Add Experience", command=lambda: add_experience(root, experience_frame, experience_list))
    add_experience_button.pack(side=tk.TOP, fill=tk.X, pady=5)

    # Education Section
    education_label = ttk.Label(main_frame, text="Education", font=('Helvetica', 14, 'bold'))
    education_label.grid(row=6, column=0, columnspan=2, sticky='w', pady=(10, 0))
    education_frame = ttk.Frame(main_frame, padding=(10, 10, 10, 10), style='TFrame')
    education_frame.grid(row=7, column=0, columnspan=2, sticky='ew')
    add_education_button = ttk.Button(education_frame, text="Add Education", command=lambda: add_education(root, education_frame, education_list))
    add_education_button.pack(side=tk.TOP, fill=tk.X, pady=5)

    # Certificates Section
    certificates_label = ttk.Label(main_frame, text="Certificates", font=('Helvetica', 14, 'bold'))
    certificates_label.grid(row=12, column=0, columnspan=2, sticky='w', pady=(10, 0))
    certificates_frame = ttk.Frame(main_frame, padding=(10, 10, 10, 10), style='TFrame')
    certificates_frame.grid(row=13, column=0, columnspan=2, sticky='ew')
    add_certificate_button = ttk.Button(certificates_frame, text="Add Certificate", command=lambda: add_certificate(root, certificates_frame, certificates_list))
    add_certificate_button.pack(side=tk.TOP, fill=tk.X, pady=5)

    # Skills Section
    skills_label = ttk.Label(main_frame, text="Skills", font=('Helvetica', 14, 'bold'))
    skills_label.grid(row=8, column=0, columnspan=2, sticky='w', pady=(10, 0))
    skills_frame = ttk.Frame(main_frame, padding=(10, 10, 10, 10), style='TFrame')
    skills_frame.grid(row=9, column=0, columnspan=2, sticky='ew')
    add_skill_button = ttk.Button(skills_frame, text="Add Skill", command=lambda: add_skill(root, skills_frame, skills_list))
    add_skill_button.pack(side=tk.TOP, fill=tk.X, pady=5)  

    # Awards Section
    awards_label = ttk.Label(main_frame, text="Awards", font=('Helvetica', 14, 'bold'))
    awards_label.grid(row=10, column=0, columnspan=2, sticky='w', pady=(10, 0))
    awards_frame = ttk.Frame(main_frame, padding=(10, 10, 10, 10), style='TFrame')
    awards_frame.grid(row=11, column=0, columnspan=2, sticky='ew')
    add_award_button = ttk.Button(awards_frame, text="Add Award", command=lambda: add_award(root, awards_frame, awards_list))
    add_award_button.pack(side=tk.TOP, fill=tk.X, pady=5)

    # Photo Upload Section
    photo_label = tk.Label(root)
    photo_label.pack(pady=10)
    button_frame = ttk.Frame(main_frame, style='TFrame')
    button_frame.grid(row=2, column=0, columnspan=2, sticky='ew')  # Adjust row number as needed
    upload_button = ttk.Button(button_frame, text="Upload Photo", command=lambda: upload_photo(photo_label))
    upload_button.grid(row=0, column=0, padx=5, pady=10)
    remove_button = ttk.Button(button_frame, text="Remove Photo", command=lambda: remove_photo(photo_label))
    remove_button.grid(row=0, column=1, padx=5, pady=10)

    # Save and Load Buttons
    buttons_frame = ttk.Frame(main_frame, style='TFrame')
    buttons_frame.grid(row=3, column=0, columnspan=2, sticky='ew')  # Adjust row number as needed
    load_button = ttk.Button(buttons_frame, text="Load CV", command=lambda: load_data(root, name_entry, email_entry, cellphone_entry, address_entry, linkedin_entry, id_number_entry, drivers_license_entry, citizenship_entry, passport_number_entry, experience_list, education_list, skills_list, awards_list, experience_frame, education_frame, skills_frame, awards_frame, certificates_frame, certificates_list))
    load_button.grid(row=0, column=0, sticky='ew', padx=5, pady=10)
    save_button = ttk.Button(buttons_frame, text="Save", command=lambda: save_data(name_entry, email_entry, cellphone_entry, address_entry, linkedin_entry, id_number_entry, drivers_license_entry, citizenship_entry, passport_number_entry, experience_list, education_list, skills_list, awards_list, certificates_list))
    save_button.grid(row=0, column=1, sticky='ew', padx=5, pady=10)
    export_button = ttk.Button(buttons_frame, text="Export to DOCX", command=lambda: export_to_docx(collect_cv_data(name_entry, email_entry, cellphone_entry, address_entry, linkedin_entry, id_number_entry, drivers_license_entry, citizenship_entry, passport_number_entry, experience_list, education_list, skills_list, awards_list, certificates_list)))
    export_button.grid(row=0, column=2, sticky='ew', padx=5, pady=10)

    root.mainloop()

def add_experience(root, experiences_frame, experience_list):
    new_experience = tk.Frame(experiences_frame)
    new_experience.pack(pady=5)

    job_title_entry = tk.Entry(new_experience)
    job_title_entry.grid(row=0, column=1, padx=10)
    tk.Label(new_experience, text="Job Title:").grid(row=0, column=0)

    company_entry = tk.Entry(new_experience)
    company_entry.grid(row=1, column=1, padx=10)
    tk.Label(new_experience, text="Company:").grid(row=1, column=0)

    period_entry = tk.Entry(new_experience)
    period_entry.grid(row=2, column=1, padx=10)
    tk.Label(new_experience, text="Period:").grid(row=2, column=0)

    description_entry = tk.Text(new_experience, height=4, width=40)  # Adjust height and width as needed
    description_entry.grid(row=3, column=1, padx=10, pady=10)
    tk.Label(new_experience, text="Description:").grid(row=3, column=0)

    remove_button = tk.Button(new_experience, text="Remove", command=lambda: remove_entry(new_experience, (job_title_entry, company_entry, period_entry, description_entry), experience_list))
    remove_button.grid(row=4, column=1, pady=5)

    experience_list.append((job_title_entry, company_entry, period_entry, description_entry))

def add_education(root, education_frame, education_list):
    new_education = tk.Frame(education_frame)
    new_education.pack(pady=5)

    school_entry = tk.Entry(new_education)
    school_entry.grid(row=0, column=1, padx=10)
    tk.Label(new_education, text="School:").grid(row=0, column=0)

    degree_entry = tk.Entry(new_education)
    degree_entry.grid(row=1, column=1, padx=10)
    tk.Label(new_education, text="Degree:").grid(row=1, column=0)

    major_entry = tk.Entry(new_education)
    major_entry.grid(row=2, column=1, padx=10)
    tk.Label(new_education, text="Major:").grid(row=2, column=0)

    dates_entry = tk.Entry(new_education)
    dates_entry.grid(row=3, column=1, padx=10)
    tk.Label(new_education, text="Dates:").grid(row=3, column=0)

    remove_button = tk.Button(new_education, text="Remove", command=lambda: remove_entry(new_education, (school_entry, degree_entry, major_entry, dates_entry), education_list))
    remove_button.grid(row=4, column=1, pady=5)

    education_list.append((school_entry, degree_entry, major_entry, dates_entry))

def add_award(root, awards_frame, awards_list):
    new_award = tk.Frame(awards_frame)
    new_award.pack(pady=5)

    award_entry = tk.Entry(new_award)
    award_entry.grid(row=0, column=1, padx=10)
    tk.Label(new_award, text="Award:").grid(row=0, column=0)

    awarded_by_entry = tk.Entry(new_award)
    awarded_by_entry.grid(row=1, column=1, padx=10)
    tk.Label(new_award, text="Awarded by:").grid(row=1, column=0)

    date_entry = tk.Entry(new_award)
    date_entry.grid(row=2, column=1, padx=10)
    tk.Label(new_award, text="Date:").grid(row=2, column=0)

    remove_button = tk.Button(new_award, text="Remove", command=lambda: remove_entry(new_award, (award_entry, awarded_by_entry, date_entry), awards_list))
    remove_button.grid(row=3, column=1, pady=5)

    awards_list.append((award_entry, awarded_by_entry, date_entry))

def add_skill(root, skills_frame, skills_list):
    new_skill = tk.Frame(skills_frame)
    new_skill.pack(pady=5)

    # Dropdown menu for skill type
    skill_types = ["Hardware", "Software"]
    skill_type_var = tk.StringVar(new_skill)
    skill_type_var.set(skill_types[0])  # default value
    skill_type_menu = tk.OptionMenu(new_skill, skill_type_var, *skill_types)
    skill_type_menu.grid(row=0, column=0, padx=10)

    skill_entry = tk.Entry(new_skill)
    skill_entry.grid(row=0, column=1, padx=10)

    remove_button = tk.Button(new_skill, text="Remove", command=lambda: remove_entry(new_skill, (skill_type_var, skill_entry), skills_list))
    remove_button.grid(row=0, column=2, padx=5)

    skills_list.append((skill_type_var, skill_entry))

def remove_entry(entry_frame, entry_tuple, entry_list):
    entry_frame.destroy()
    entry_list.remove(entry_tuple)

def add_certificate(root, certificates_frame, certificates_list):
    new_certificate = tk.Frame(certificates_frame)
    new_certificate.pack(pady=5)

    cert_title_entry = tk.Entry(new_certificate)
    cert_title_entry.grid(row=0, column=1, padx=10)
    tk.Label(new_certificate, text="Certificate Title:").grid(row=0, column=0)

    issued_by_entry = tk.Entry(new_certificate)
    issued_by_entry.grid(row=1, column=1, padx=10)
    tk.Label(new_certificate, text="Issued by:").grid(row=1, column=0)

    date_entry = tk.Entry(new_certificate)
    date_entry.grid(row=2, column=1, padx=10)
    tk.Label(new_certificate, text="Date:").grid(row=2, column=0)

    remove_button = tk.Button(new_certificate, text="Remove", command=lambda: remove_entry(new_certificate, (cert_title_entry, issued_by_entry, date_entry), certificates_list))
    remove_button.grid(row=3, column=1, pady=5)

    certificates_list.append((cert_title_entry, issued_by_entry, date_entry))

def save_data(name_entry, email_entry, cellphone_entry, address_entry, linkedin_entry, id_number_entry, drivers_license_entry, citizenship_entry, passport_number_entry, experience_list, education_list, skills_list, awards_list, certificates_list):
    try:
        name = name_entry.get()
        email = email_entry.get()
        cellphone = cellphone_entry.get()
        address = address_entry.get().strip()
        linkedin = linkedin_entry.get()
        id_number = id_number_entry.get().strip()
        drivers_license = drivers_license_entry.get().strip()
        citizenship = citizenship_entry.get().strip()
        passport_number = passport_number_entry.get().strip()

        cv_data = {
            "personal_info": {
                "name": name_entry.get(),
                "email": email_entry.get(),
                "cellphone": cellphone_entry.get(),
                "address": address_entry.get().strip(),
                "linkedin": linkedin_entry.get(),
                "id_number": id_number_entry.get(),
                "drivers_license": drivers_license_entry.get(),
                "citizenship": citizenship_entry.get(),
                "passport_number": passport_number_entry.get()
            },
            "experience": [],
            "education": [],
            "skills": [],
            "awards": [],
            "certificates": []
        }
    
        # Validate personal information
        errors = []
        if not name.strip():
            errors.append("Name cannot be empty.")
        if not is_valid_email(email):
            errors.append("Invalid email format.")
        if not len(cellphone.strip()) <= 12:
            errors.append("Cellphone number must be 10 digits.")

        # Check if there are any errors
        if errors:
            error_message = "\n".join(errors)
            messagebox.showerror("Validation Error", error_message)
            return  # Stop the save process if there are validation errors

        for exp in experience_list:
            cv_data["experience"].append({
                "job_title": exp[0].get(),
                "company": exp[1].get(),
                "period": exp[2].get(),
                "description": exp[3].get('1.0', 'end-1c')
            })

        for edu in education_list:
            cv_data["education"].append({
                "school": edu[0].get(),
                "degree": edu[1].get(),
                "major": edu[2].get(),
                "dates": edu[3].get()
            })

        for cert in certificates_list:
            cv_data["certificates"].append({
                "title": cert[0].get(),
                "issued_by": cert[1].get(),
                "date": cert[2].get()
            })

        cv_data["skills"] = [{"type": skill[0].get(), "description": skill[1].get()} for skill in skills_list]

        for award in awards_list:
            cv_data["awards"].append({
                "award": award[0].get(),
                "awarded_by": award[1].get(),
                "date": award[2].get()
            })

        try:
            with open("cv_data.json", "w") as file:
                json.dump(cv_data, file, indent=4)
        except IOError as e:
            messagebox.showerror("Save Error", f"An error occurred while saving: {e}")

    except IOError as e:
        messagebox.showerror("Save Error", f"An error occurred while saving: {e}")
    except Exception as e:
        messagebox.showerror("Save Error", f"An unexpected error occurred: {e}")
        print(traceback.format_exc())

def load_data(root, name_entry, email_entry, cellphone_entry, address_entry, linkedin_entry, id_number_entry, drivers_license_entry, citizenship_entry, passport_number_entry, experience_list, education_list, skills_list, awards_list, experiences_frame, education_frame, skills_frame, awards_frame, certificates_frame, certificates_list):
    try:
        with open("cv_data.json", "r") as file:
            data = json.load(file)

        # Load personal information
        personal_info = data.get("personal_info", {})
        name_entry.delete(0, tk.END)
        name_entry.insert(0, personal_info.get("name", ""))
        email_entry.delete(0, tk.END)
        email_entry.insert(0, personal_info.get("email", ""))
        cellphone_entry.delete(0, tk.END)
        cellphone_entry.insert(0, personal_info.get("cellphone", ""))
        address_entry.delete(0, tk.END)
        address_entry.insert(0, personal_info.get("address", ""))
        linkedin_entry.delete(0, tk.END)
        linkedin_entry.insert(0, personal_info.get("linkedin", ""))
        id_number_entry.delete(0, tk.END)
        id_number_entry.insert(0, personal_info.get("id_number", ""))
        drivers_license_entry.delete(0, tk.END)
        drivers_license_entry.insert(0, personal_info.get("drivers_license", ""))
        citizenship_entry.delete(0, tk.END)
        citizenship_entry.insert(0, personal_info.get("citizenship", ""))
        passport_number_entry.delete(0, tk.END)
        passport_number_entry.insert(0, personal_info.get("passport_number", ""))

        # Check and load experiences
        if "experience" in data:
            for exp in data["experience"]:
                add_experience(root, experiences_frame, experience_list)
                experience_list[-1][0].insert(0, exp["job_title"])
                experience_list[-1][1].insert(0, exp["company"])
                experience_list[-1][2].insert(0, exp["period"])
                experience_list[-1][3].insert(1.0, exp["description"])

        # Load education
        if "education" in data:
            for edu in data["education"]:
                add_education(root, education_frame, education_list)
                education_list[-1][0].insert(0, edu["school"])
                education_list[-1][1].insert(0, edu["degree"])
                education_list[-1][2].insert(0, edu["major"])
                education_list[-1][3].insert(0, edu["dates"])

        # Load skills
        if "skills" in data:
            for skill in data["skills"]:
                add_skill(root, skills_frame, skills_list)
                skills_list[-1][0].set(skill["type"])  # Set the skill type
                skills_list[-1][1].insert(0, skill["description"])  # Set the skill description

        # Load awards
        if "awards" in data:
            for award in data["awards"]:
                add_award(root, awards_frame, awards_list)
                awards_list[-1][0].insert(0, award["award"])
                awards_list[-1][1].insert(0, award["awarded_by"])
                awards_list[-1][2].insert(0, award["date"])

        if "certificates" in data:
            for cert in data["certificates"]:
                add_certificate(root, certificates_frame, certificates_list)
                certificates_list[-1][0].insert(0, cert["title"])
                certificates_list[-1][1].insert(0, cert["issued_by"])
                certificates_list[-1][2].insert(0, cert["date"])

    except FileNotFoundError:
        tk.messagebox.showerror("Load Error", "The CV data file could not be found.")
    except json.JSONDecodeError:
        tk.messagebox.showerror("Load Error", "The CV data file is not in the correct JSON format.")
    except Exception as e:
        tk.messagebox.showerror("Load Error", f"An unexpected error occurred: {e}")
        print(traceback.format_exc())

def export_to_docx(cv_data):
    try:
        doc = Document()
        doc.add_heading('Curriculum Vitae', 0)

        if photo_path:
            doc.add_picture(photo_path, width=Pt(100), height=Pt(100))

        # Personal Info
        personal_info = cv_data["personal_info"]
        doc.add_heading('Personal Information', level=1)
        doc.add_paragraph(f"Name: {personal_info['name']}")
        doc.add_paragraph(f"Email: {personal_info['email']}")
        doc.add_paragraph(f"Cellphone: {personal_info['cellphone']}")
        doc.add_paragraph(f"LinkedIn: {personal_info['linkedin']}")
        doc.add_paragraph(f"ID Number: {personal_info['id_number']}")
        doc.add_paragraph(f"Drivers License Code: {personal_info['drivers_license']}")
        doc.add_paragraph(f"Citizenship: {personal_info['citizenship']}")
        doc.add_paragraph(f"Passport Number: {personal_info['passport_number']}")
        # Add other personal info fields as needed

        # Experience
        doc.add_heading('Experience', level=1)
        for exp in cv_data["experience"]:
            p = doc.add_paragraph()
            p.add_run(exp['job_title']).bold = True
            p.add_run(f" at {exp['company']} - {exp['period']}\n").italic = True
            p.add_run(exp['description'])

        # Education
        doc.add_heading('Education', level=1)
        for edu in cv_data["education"]:
            p = doc.add_paragraph()
            p.add_run(edu['degree']).bold = True
            p.add_run(f" at {edu['school']} - {edu['dates']}\n").italic = True
            p.add_run(edu['major'])

        # Certificates
        doc.add_heading('Certificates', level=1)
        for cert in cv_data["certificates"]:
            doc.add_paragraph(f"{cert['title']} - Issued by {cert['issued_by']} on {cert['date']}")

        # Skills
        doc.add_heading('Skills', level=1)
        for skill in cv_data["skills"]:
            doc.add_paragraph(f"{skill['type']}: {skill['description']}")

        # Awards
        doc.add_heading('Awards', level=1)
        for award in cv_data["awards"]:
            p = doc.add_paragraph()
            p.add_run(award['award']).bold = True
            p.add_run(f" by {award['awarded_by']} - {award['date']}")

        # Save the document
        doc.save('CV.docx')
    
    except Exception as e:
        messagebox.showerror("Export Error", f"An unexpected error occurred: {e}")
        print(traceback.format_exc())

def collect_cv_data(name_entry, email_entry, cellphone_entry, address_entry, linkedin_entry, id_number_entry, drivers_license_entry, citizenship_entry, passport_number_entry, experience_list, education_list, skills_list, awards_list, certificates_list):
    # Collect Personal Information
    personal_info = {
        "name": name_entry.get(),
        "email": email_entry.get(),
        "cellphone": cellphone_entry.get(),
        "linkedin": linkedin_entry.get(),
        "address": address_entry.get().strip(),
        "id_number": id_number_entry.get().strip(),
        "drivers_license": drivers_license_entry.get().strip(),
        "citizenship": citizenship_entry.get().strip(),
        "passport_number": passport_number_entry.get().strip()
    }

    # Collect Experiences
    experiences = []
    for job_title_entry, company_entry, period_entry, description_entry in experience_list:
        experiences.append({
            "job_title": job_title_entry.get(),
            "company": company_entry.get(),
            "period": period_entry.get(),
            "description": description_entry.get("1.0", tk.END).strip()
        })

    # Collect Education
    educations = []
    for school_entry, degree_entry, major_entry, dates_entry in education_list:
        educations.append({
            "school": school_entry.get(),
            "degree": degree_entry.get(),
            "major": major_entry.get(),
            "dates": dates_entry.get()
        })

    # Collect Skills
    skills = []
    for skill_type_var, skill_entry in skills_list:
        skills.append({
            "type": skill_type_var.get(),
            "description": skill_entry.get()
        })

    # Collect Awards
    awards = []
    for award_entry, awarded_by_entry, date_entry in awards_list:
        awards.append({
            "award": award_entry.get(),
            "awarded_by": awarded_by_entry.get(),
            "date": date_entry.get()
        })

    # Collect Certificates
    certificates = []
    for cert_title_entry, issued_by_entry, date_entry in certificates_list:
        certificates.append({
            "title": cert_title_entry.get(),
            "issued_by": issued_by_entry.get(),
            "date": date_entry.get()
        })

    # Structure the complete CV data
    cv_data = {
        "personal_info": personal_info,
        "experience": experiences,
        "education": educations,
        "skills": skills,
        "awards": awards,
        "certificates": certificates
    }

    return cv_data

if __name__ == "__main__":
    main()